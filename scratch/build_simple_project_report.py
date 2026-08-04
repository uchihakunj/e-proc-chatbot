from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "docs" / "reports" / "CHiPS_eProc_Chatbot_Simple_Project_Report.docx"
ASSETS = ROOT / "docs" / "reports" / "monthly_work_done_assets"

FONT = "Arial"
SIZE = Pt(12)
BLUE = RGBColor(31, 78, 121)
BLACK = RGBColor(0, 0, 0)
LIGHT_BLUE = "D9EAF7"


def set_font(run, bold=None, color=BLACK, italic=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = SIZE
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    props.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_field(paragraph):
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph.add_run()._r.append(field)


def set_paragraph_format(paragraph, before=0, after=6, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def write_paragraph(doc, text="", bold=False, color=BLACK, align=None, before=0, after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    set_paragraph_format(p, before, after)
    r = p.add_run(text)
    set_font(r, bold=bold, color=color)
    return p


def write_heading(doc, text):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=12, after=6)
    r = p.add_run(text)
    set_font(r, bold=True, color=BLUE)
    return p


def write_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_format(p, before=0, after=3)
    r = p.add_run(text)
    set_font(r)
    return p


def write_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_paragraph_format(p, before=0, after=3)
    r = p.add_run(text)
    set_font(r)
    return p


def style_cell(cell, text, bold=False, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_format(p, before=0, after=0, line=1.1)
    r = p.add_run(text)
    set_font(r, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    if fill:
        shade(cell, fill)


def setup_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = SIZE
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name in ("Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number", "Header", "Footer"):
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = SIZE

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_format(header, 0, 0, 1.0)
    set_font(header.add_run("CHiPS e-Procurement AI Chatbot | Project Report"), bold=True, color=BLUE)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(footer, 0, 0, 1.0)
    set_font(footer.add_run("Page "))
    add_page_field(footer)


def add_technology_table(doc):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(2.25)
    table.columns[1].width = Inches(4.65)
    style_cell(table.rows[0].cells[0], "Technology", bold=True, fill=LIGHT_BLUE)
    style_cell(table.rows[0].cells[1], "Use in the Project", bold=True, fill=LIGHT_BLUE)
    rows = [
        ("Python and Flask", "Backend API, query processing, response streaming, and application orchestration."),
        ("HTML, CSS, JavaScript", "Responsive chatbot interface, quick prompts, answer display, and user interaction."),
        ("Node.js / Express", "Web UI serving and proxy support for the browser-facing application."),
        ("Qdrant", "Vector database for storing and searching the processed document knowledge base."),
        ("BGE-M3 and reranker", "Semantic embedding and re-ranking of relevant procurement document passages."),
        ("Ollama-hosted models", "Local language-model inference for conversational, context-aware answers."),
        ("OCR and document tools", "Preparation, extraction, cleanup, and chunking of source PDF manuals and rules."),
        ("Voice interaction tools", "Optional speech-based interaction to improve accessibility and usability."),
    ]
    for tech, purpose in rows:
        cells = table.add_row().cells
        style_cell(cells[0], tech, bold=True)
        style_cell(cells[1], purpose)
    return table


def add_screenshot(doc, caption, path):
    p = write_paragraph(doc, caption, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=10, after=4)
    doc.add_picture(str(path), width=Inches(6.5))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(last, before=0, after=10)


def build_report():
    doc = Document()
    setup_document(doc)

    write_paragraph(
        doc,
        "CHiPS e-Procurement AI Chatbot Project Report",
        bold=True,
        color=BLUE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=18,
        after=18,
    )

    write_heading(doc, "Title")
    write_paragraph(doc, "CHiPS e-Procurement AI Chatbot")

    write_heading(doc, "Problem Statement")
    write_paragraph(
        doc,
        "Government e-procurement users must often search through lengthy manuals, circulars, rules, and portal instructions to complete tasks such as vendor registration, bid submission, tender participation, EMD handling, and auction-related activities. This process can be slow and difficult, particularly when information is spread across multiple documents or when users need support in Hindi, English, or Hinglish."
    )
    write_paragraph(
        doc,
        "The project addresses this challenge by providing a simple conversational interface that gives document-grounded guidance while retaining a clear connection to the underlying source material."
    )

    write_heading(doc, "Approach")
    write_paragraph(doc, "The solution follows a Retrieval-Augmented Generation (RAG) approach:")
    for item in [
        "Collect procurement manuals, rules, FAQs, and portal guidance documents.",
        "Prepare the documents through image cleanup, OCR where required, text extraction, and chunking.",
        "Convert the content into embeddings and store it in a searchable vector database.",
        "Retrieve and re-rank the most relevant source passages for every user question.",
        "Generate a concise answer from the retrieved context and present it through the web chatbot.",
        "Apply routing, evidence, and language controls to keep answers relevant, source-grounded, and suitable for English, Hindi, and Hinglish queries.",
    ]:
        write_numbered(doc, item)

    write_heading(doc, "Technologies Used")
    add_technology_table(doc)

    write_heading(doc, "Today's Migration Brief - GeM Catalogue Integration")
    write_paragraph(
        doc,
        "A safeguarded catalogue capability has been added for factual product-oriented queries while procurement-policy questions continue through the existing RAG workflow."
    )

    write_paragraph(doc, "Completed Today", bold=True, color=BLUE, before=4, after=4)
    for item in [
        "Implemented an isolated periodic GeM catalogue scraper and a public GeM product/search-page parser.",
        "Added the PostgreSQL catalogue schema and Alembic migration.",
        "Added incremental synchronization using content hashes and stale-cache detection.",
        "Added robots, HTTPS, host, rate-limit, and CAPTCHA safeguards.",
        "Added cached catalogue lookup routing for price, stock, availability, and product listings.",
        "Added catalogue support in POST /api/catalogue/search, GET /api/catalogue/status, /api/query, and /api/stream.",
        "Added test-only CATALOGUE_ONLY_TEST_MODE while preserving existing backend procurement behavior.",
    ]:
        write_bullet(doc, item)

    write_paragraph(doc, "Verification Completed", bold=True, color=BLUE, before=8, after=4)
    for item in [
        "Catalogue and scraper tests passed.",
        "API integration tests passed.",
        "Existing chatbot regression tests passed.",
        "Total verified tests: 36 of 36 passed.",
        "The temporary sample database was removed after verification.",
    ]:
        write_bullet(doc, item)

    write_paragraph(doc, "Deployment Status and Next Steps", bold=True, color=BLUE, before=8, after=4)
    for item in [
        "Create the PostgreSQL role and gem_catalog database, then add the DATABASE_URL and catalogue environment settings. Use the sample provider only for database testing.",
        "Run the Alembic migration, insert sample catalogue records, restart the backend, and check /api/catalogue/status for a healthy state.",
        "Before production use, configure the approved public GeM search URL, probe the source, perform the first real synchronization, enable the periodic timer, and test current-price and availability questions in the web UI.",
        "Real GeM data has not yet been loaded, and the running backend is still the old process until it is restarted. Catalogue answers must not be presented as live GeM data until the real synchronization succeeds.",
    ]:
        write_bullet(doc, item)

    write_heading(doc, "Screenshots")
    write_paragraph(doc, "The following screens show the chatbot’s portal presentation, open interaction view, and a sample answer state.")
    add_screenshot(doc, "Figure 1. Portal landing page with chatbot entry point", ASSETS / "chatbot_home.png")
    add_screenshot(doc, "Figure 2. Open chatbot widget for user interaction", ASSETS / "chatbot_widget_open.png")
    add_screenshot(doc, "Figure 3. Sample source-grounded chatbot response", ASSETS / "chatbot_sample_answer.png")

    write_heading(doc, "Project Outcome")
    write_paragraph(
        doc,
        "The CHiPS e-Procurement AI Chatbot brings procurement information into a more accessible, multilingual, and user-friendly format. It combines a document-processing pipeline, semantic retrieval, local language-model generation, a web interface, and optional voice support to help users obtain guidance quickly while maintaining source awareness."
    )

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_report()
