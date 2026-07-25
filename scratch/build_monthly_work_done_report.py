from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "reports"
ASSET_DIR = OUT_DIR / "monthly_work_done_assets"
OUT_PATH = OUT_DIR / "Monthly_Work_Done_Report_July_2026.docx"

BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(31, 41, 55)
LIGHT = RGBColor(242, 246, 252)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_page_number(paragraph):
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    pf = normal.paragraph_format
    pf.space_after = Pt(4)
    pf.line_spacing = 1.12

    for name, size in [("Heading 1", 14), ("Heading 2", 11.5)]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("CHiPS e-Procurement AI Chatbot")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = BLUE

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Monthly Work Done Report | Page ")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    set_page_number(footer)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Submission Requirements\nMonthly Work Done Report")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(20)
    run.font.color.rgb = BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Project: CHiPS e-Procurement AI Chatbot\nReporting Period: 01 July 2026 to 24 July 2026")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = DARK

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.style = "Table Grid"
    meta.autofit = False
    widths = [Inches(2.2), Inches(4.9)]
    rows = [
        ("Prepared for", "Project Submission / Monthly Progress Review"),
        ("Prepared by", "Development Team"),
        ("Project Scope", "Multilingual RAG chatbot for e-Procurement manuals, rules, bids, tenders, EMD, and vendor workflows"),
        ("Current Focus", "Answer quality improvement, routing controls, UI polish, voice support, and validation"),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], "D9EAF7")
        for p in row.cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Arial"
                r.font.size = Pt(10)
        for p in row.cells[1].paragraphs:
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(10)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        r = p.add_run(f"• {item}")
        r.font.name = "Arial"
        r.font.size = Pt(10.5)


def add_sections(doc):
    doc.add_paragraph()
    doc.add_paragraph("1. Executive Summary", style="Heading 1")
    doc.add_paragraph(
        "During July 2026, the CHiPS e-Procurement AI Chatbot moved from a broad diagnostic phase into targeted quality improvement and presentation readiness. "
        "Work focused on strengthening source-grounded responses, improving actor and intent routing, tightening policy-sensitive answers, enhancing the user interface, "
        "and adding voice and browser-friendly interaction paths."
    )
    doc.add_paragraph(
        "The month also included structured validation through focused live API checks, UAT answer-quality audits, and a production benchmark harness covering multilingual procurement queries."
    )

    doc.add_paragraph("2. Major Work Completed", style="Heading 1")
    doc.add_paragraph("2.1 UI and User Experience Enhancements", style="Heading 2")
    add_bullets(
        doc,
        [
            "Refined the public chatbot widget and embedded it within a government-style portal background for a more production-ready user experience.",
            "Added quick-question chips, welcome prompts, source drawer, highlighted PDF viewing flow, and a cleaner answer presentation layer.",
            "Enabled open-access chat flow with simplified user journey and improved visual hierarchy for English, Hindi, and Hinglish usage.",
            "Added Auto-Voice support and improved voice interaction handling for easier hands-free use.",
        ],
    )

    doc.add_paragraph("2.2 RAG Quality and Answer Controls", style="Heading 2")
    add_bullets(
        doc,
        [
            "Improved actor routing, fine-intent controls, context diagnostics, and response-language enforcement to reduce answer drift.",
            "Added deterministic and source-bound answers for high-risk procurement topics such as purchase splitting, DSC obtainment, bid eligibility, technical versus financial bid handling, and evaluation reporting.",
            "Strengthened policy-sensitive responses to avoid unsupported portal claims and to keep answers aligned with supplied manuals and rules.",
            "Introduced answer-quality scoring in the benchmark flow to measure narrow-question accuracy instead of only broad factual overlap.",
        ],
    )

    doc.add_paragraph("2.3 Validation, Audit, and Engineering Support", style="Heading 2")
    add_bullets(
        doc,
        [
            "Ran focused live API validation and regression testing after each major answer-policy change.",
            "Completed repository organization updates, deployment support files, and maintenance utilities to make the project easier to operate and review.",
            "Added browser STT fallback work and voice latency improvements to support practical user interaction in the web interface.",
        ],
    )

    doc.add_paragraph("3. Validation Highlights", style="Heading 1")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    headers = ["Validation Item", "Result", "Observation"]
    widths = [Inches(2.3), Inches(1.3), Inches(3.5)]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = text
        cell.width = widths[i]
        set_cell_shading(cell, "D9EAF7")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Arial"
                r.font.size = Pt(10)

    data = [
        ("Exact-answer regression suite", "49 passed", "Validated actor boundaries, narrow-answer contracts, and procurement workflow cases."),
        ("Focused live API check", "12 / 12 pass", "Mean latency 1.51 seconds; maximum 1.82 seconds for targeted production questions."),
        ("UAT quality remediation tests", "37 passed", "Confirmed fixes for purchase splitting, eligibility, DSC, and report-generation answer behavior."),
        ("Production benchmark harness", "120 queries executed", "Used as a defect-discovery baseline for multilingual routing, retrieval, citation, and answer completeness."),
    ]
    for row_data in data:
        row = table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.width = widths[i]
            cell.text = text
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if i == 1:
                set_cell_shading(cell, "F3F8FC")

    doc.add_paragraph("4. Key Outcomes for the Month", style="Heading 1")
    add_bullets(
        doc,
        [
            "Higher confidence in source-grounded chatbot responses for procurement procedures and compliance-sensitive queries.",
            "More polished and user-friendly chatbot experience with embedded UI, quick prompts, source exploration, and voice controls.",
            "Improved engineering readiness through stronger diagnostics, maintenance scripts, and clearer repository structure.",
            "A repeatable benchmarking and regression workflow that will support the next round of production improvements.",
        ],
    )

    doc.add_paragraph("5. Work Planned for the Next Cycle", style="Heading 1")
    add_bullets(
        doc,
        [
            "Re-run the expanded multilingual UAT after the latest routing and answer-contract changes.",
            "Continue reducing fallback behavior and improving answer completeness for department-operator and vendor workflows.",
            "Extend UI verification in a live browser environment and continue screenshot-backed product documentation.",
            "Prepare the chatbot stack for broader deployment review and stakeholder demonstrations.",
        ],
    )


def add_screenshots(doc):
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    section = doc.sections[-1]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    doc.add_paragraph("6. Chatbot Screenshots", style="Heading 1")
    doc.add_paragraph(
        "The following screenshots capture the current chatbot presentation layer used for demonstration and monthly submission purposes."
    )

    shots = [
        ("A. Portal landing page with AI entry point", ASSET_DIR / "chatbot_home.png", 6.8),
        ("B. Chatbot widget opened from the portal", ASSET_DIR / "chatbot_widget_open.png", 6.8),
        ("C. Sample FAQ response shown inside the chatbot", ASSET_DIR / "chatbot_sample_answer.png", 6.8),
    ]
    for title, path, width in shots:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(10.5)
        doc.add_picture(str(path), width=Inches(width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(path.name.replace("_", " ").replace(".png", "").title())
        run.italic = True
        run.font.name = "Arial"
        run.font.size = Pt(9)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_doc(doc)
    add_header_footer(doc)
    add_title(doc)
    doc.add_paragraph()
    add_sections(doc)
    add_screenshots(doc)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
