from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT = Path(__file__).resolve().parents[1] / "docs" / "reports" / "Chhattisgarh_eProcurement_Chatbot_Project_Report.docx"


def set_font(run, size, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def set_cell_border(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), "000000")


def body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, 12)
    return p


def heading(doc, text, size=16):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    set_font(r, size, True)
    return p


def screenshot_block(doc, number, name, purpose, bullets, image_path):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    set_font(p.add_run(f"Screenshot {number} - {name}"), 14, True)
    body_paragraph(doc, f"Purpose: {purpose}")
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    set_font(p.add_run("Description"), 12, True)
    for item in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(2)
        set_font(p.add_run(item), 12)
    image = doc.add_picture(str(image_path), width=Inches(4.55))
    image.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(12)
    set_font(cap.add_run(f"Figure {number} - {name}"), 10)


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)
    set_font(title.add_run("Chhattisgarh e-Procurement Chatbot"), 18, True)

    heading(doc, "Module")
    body_paragraph(doc, "Procurement Policy Query and Guidance System")

    heading(doc, "Team Members")
    p = doc.add_paragraph(style="List Bullet")
    set_font(p.add_run("[Information to be provided]"), 12)

    heading(doc, "Problem Statement")
    body_paragraph(doc, "Procurement personnel need timely and consistent guidance on rules, procedures, approvals, and document requirements. Relevant information is distributed across policy documents, circulars, manuals, and departmental references. Searching these sources manually is time-consuming and can lead to incomplete or inconsistent responses. The issue affects officials who require clear guidance while preparing procurement requests, processing purchases, and responding to routine procedural questions. A single question may require reviewing multiple documents, which increases the effort needed for day-to-day work. The system is intended to improve access to approved procurement information while retaining the original source material as the basis for responses.")

    heading(doc, "Solution Proposed")
    body_paragraph(doc, "The Chhattisgarh e-Procurement Chatbot has been developed as a question-answering system for procurement-policy guidance. It receives a user query, identifies the relevant procurement context, retrieves supporting content from the approved knowledge base, and presents a clear response through a web interface. The system supports natural-language questions and is designed to provide concise guidance for common procurement processes. It maintains separation between source retrieval, response generation, and the user interface so that each component can be maintained independently. The expected outcome is faster access to procurement guidance, reduced manual search effort, and more consistent handling of routine questions by procurement personnel.")

    heading(doc, "Approach")
    steps = [
        "Input: The user submits a procurement-related question through the web interface.",
        "Processing: The system identifies the query context and searches relevant approved content.",
        "Data Storage: Source documents, processed text, and retrieval data are maintained separately.",
        "AI/ML Processing: Retrieval-augmented generation prepares a response using relevant source content.",
        "Output: The chatbot returns a clear procurement-policy response to the user.",
    ]
    for item in steps:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)
        set_font(p.add_run(item), 12)

    heading(doc, "Technology Used")
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.style = "Table Grid"
    table.columns[0].width = Inches(2)
    table.columns[1].width = Inches(4.5)
    headers = ["Component", "Technology"]
    for cell, value in zip(table.rows[0].cells, headers):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.width = Inches(2 if value == "Component" else 4.5)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_font(p.add_run(value), 12, True)
    rows = [
        ("Frontend", "HTML5, CSS3, JavaScript, Node.js and Express.js"),
        ("Backend", "Python 3.11, Flask 3 and Waitress WSGI server"),
        ("Database", "Qdrant vector database with local persistent storage"),
        ("AI/ML", "BGE-M3 embeddings, BGE reranker, Ollama, llama3.2 and llama3:8b"),
        ("APIs", "Flask REST JSON APIs and Server-Sent Events on POST /api/stream"),
        ("Frameworks", "Qdrant Client, FlagEmbedding, SentenceTransformers and Transformers"),
        ("Deployment", "Windows environment with Node.js proxy and local Python backend"),
        ("Other Tools", "PyMuPDF, OpenCV, Pillow, python-docx, Git and VS Code"),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        for cell, value, width in zip(cells, (left, right), (2, 4.5)):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.width = Inches(width)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_font(p.add_run(value), 12)

    heading(doc, "Relevant Screenshots")
    assets = OUT.parent / "assets"
    screenshot_block(doc, 1, "Portal Home Page", "Shows the public Chhattisgarh e-Procurement portal and chatbot entry point.", ["Government of Chhattisgarh header", "e-Procurement system title", "Procurement service categories", "Ask E-proc AI entry button"], assets / "chatbot_home.png")
    screenshot_block(doc, 2, "Assistant Welcome Interface", "Shows the assistant opening panel and available procurement guidance topics.", ["Assistant welcome message", "User role starting points", "Suggested procurement questions", "Question input field"], assets / "chatbot_assistant_welcome.png")
    screenshot_block(doc, 3, "Source-Backed Chatbot Response", "Shows a chatbot response with the supporting source reference.", ["User procurement question", "Generated answer", "CHiPS FAQ source reference", "Official-document verification notice"], assets / "chatbot_query_response.png")

    doc.core_properties.title = "Chhattisgarh e-Procurement Chatbot"
    doc.core_properties.author = ""
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
