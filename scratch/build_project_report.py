from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "reports"
ASSET_DIR = OUT_DIR / "monthly_work_done_assets"
OUT_PATH = OUT_DIR / "CHiPS_eProc_Chatbot_Project_Report.docx"

BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(31, 41, 55)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_page_field(paragraph):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph.add_run()._r.append(fld)


def set_base_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(4)

    for style_name, size in [("Heading 1", 14), ("Heading 2", 11.5)]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)


def add_header_footer(doc):
    header = doc.sections[0].header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("CHiPS e-Procurement AI Chatbot Project Report")
    r.font.name = "Arial"
    r.font.size = Pt(9)
    r.font.color.rgb = BLUE

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Project Report | Page ")
    r.font.name = "Arial"
    r.font.size = Pt(9)
    add_page_field(footer)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CHiPS e-Procurement AI Chatbot\nProject Report")
    r.font.name = "Arial"
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared on 24 July 2026")
    r.font.name = "Arial"
    r.font.size = Pt(11)

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(2.25), Inches(4.85)]
    data = [
        ("Project Name", "CHiPS e-Procurement AI Chatbot"),
        ("Domain", "Government e-Procurement assistance, document-grounded query answering, and multilingual support"),
        ("Supported Languages", "English, Hindi, and Hinglish"),
        ("Current Stage", "Validated prototype with web UI, source-grounded answers, voice support, and benchmark workflow"),
    ]
    for row, (k, v) in zip(table.rows, data):
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[0].text = k
        row.cells[1].text = v
        shade(row.cells[0], "D9EAF7")
        for p in row.cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Arial"
                r.font.size = Pt(10)
        for p in row.cells[1].paragraphs:
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(10)


def bullet_list(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        r = p.add_run(f"• {item}")
        r.font.name = "Arial"
        r.font.size = Pt(10.5)


def add_main_sections(doc):
    doc.add_paragraph()
    doc.add_paragraph("1. Project Goal and Key Features", style="Heading 1")
    doc.add_paragraph(
        "The CHiPS e-Procurement AI Chatbot is designed to help users understand procurement rules, portal workflows, and document-based procedures through a conversational interface. "
        "Its main objective is to provide fast, source-grounded guidance for tenders, bids, EMD, vendor registration, auctions, and procurement rules without forcing users to manually search through large manuals."
    )
    bullet_list(
        doc,
        [
            "Multilingual question answering in English, Hindi, and Hinglish.",
            "Document-grounded responses using a Retrieval-Augmented Generation workflow.",
            "Support for procurement topics such as tender methods, bid submission, vendor registration, EMD, and GFR / Store Purchase rules.",
            "Web-based chatbot widget integrated into a government-style portal interface.",
            "Source viewing support so users can connect responses back to the original manuals and reference material.",
            "Voice-oriented interaction support including Auto-Voice and speech-flow enhancements.",
        ],
    )

    doc.add_paragraph("1.1 Chatbot Interface Snapshots", style="Heading 2")
    doc.add_paragraph(
        "The screenshots below show the current chatbot presentation layer, including the portal landing view, the opened assistant widget, and a sample answer state."
    )

    screenshots = [
        ("Portal landing page with chatbot entry point", ASSET_DIR / "chatbot_home.png"),
        ("Opened chatbot widget for user interaction", ASSET_DIR / "chatbot_widget_open.png"),
        ("Sample chatbot response with answer and source", ASSET_DIR / "chatbot_sample_answer.png"),
    ]
    for caption, image_path in screenshots:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption)
        cr.bold = True
        cr.font.name = "Arial"
        cr.font.size = Pt(10.5)
        doc.add_picture(str(image_path), width=Inches(6.7))
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(image_path.name)
        cr.italic = True
        cr.font.name = "Arial"
        cr.font.size = Pt(9)

    doc.add_paragraph("2. System Design and Technology Stack", style="Heading 1")
    doc.add_paragraph(
        "The system follows a layered architecture where documents are processed into searchable knowledge assets, retrieved context is ranked and selected, and the final answer is generated and displayed in a browser-based interface."
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(1.9), Inches(2.0), Inches(2.85)]
    headers = ["Layer", "Technology", "Purpose"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        cell.text = h
        shade(cell, "D9EAF7")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Arial"
                r.font.size = Pt(10)
    rows = [
        ("Frontend UI", "HTML, CSS, JavaScript, Node.js / Express", "Provides the chatbot widget, portal-like interface, quick prompts, and UI-level interactions."),
        ("Backend API", "Python 3.11, Flask", "Handles query processing, streaming responses, routing, and answer orchestration."),
        ("Retrieval Layer", "Qdrant, BGE-M3 embeddings, bge-reranker-v2-m3", "Stores and ranks relevant document chunks for source-grounded answers."),
        ("LLM Layer", "Ollama-hosted local models", "Generates conversational answers using retrieved procurement context."),
        ("Source Corpus", "Tender manuals, GFR, IT Act, vendor and bid manuals", "Acts as the factual base for procurement guidance and portal instructions."),
        ("Voice Support", "Browser-side voice interaction and Auto-Voice flow", "Improves accessibility and ease of use for spoken interaction."),
    ]
    for row_data in rows:
        row = table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.width = widths[i]
            cell.text = text
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph("3. Week-by-Week Progress", style="Heading 1")
    week_table = doc.add_table(rows=1, cols=2)
    week_table.style = "Table Grid"
    week_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    week_table.autofit = False
    week_table.rows[0].cells[0].width = Inches(1.8)
    week_table.rows[0].cells[1].width = Inches(5.0)
    week_table.rows[0].cells[0].text = "Week"
    week_table.rows[0].cells[1].text = "Progress Achieved"
    for cell in week_table.rows[0].cells:
        shade(cell, "D9EAF7")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Arial"
                r.font.size = Pt(10)
    weeks = [
        ("Week 1", "Stabilized the project structure, aligned documentation layout, and maintained the core chatbot codebase for easier development and review."),
        ("Week 2", "Enhanced the UI with improved chatbot presentation, quick prompts, cleaner answer rendering, and Auto-Voice controls."),
        ("Week 3", "Improved actor routing, fine-intent handling, context diagnostics, and language enforcement to make answers more focused and source-bound."),
        ("Week 4", "Ran exact-answer validation, UAT answer-quality audit, and production benchmark checks; documented issues, fixes, and current performance status."),
    ]
    for week, progress in weeks:
        row = week_table.add_row()
        row.cells[0].text = week
        row.cells[1].text = progress
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph("4. Overall Summary", style="Heading 1")
    doc.add_paragraph(
        "This reporting period strengthened both the technical quality and the presentation quality of the CHiPS e-Procurement AI Chatbot. "
        "The project now has a more polished web interface, stronger answer controls for sensitive procurement topics, improved multilingual handling, and a clearer validation workflow through regression tests, live checks, and benchmark reporting. "
        "The chatbot is better positioned for demonstrations, further evaluation, and the next cycle of production-oriented refinement."
    )

    doc.add_paragraph("5. Planned Work for the Next Period", style="Heading 1")
    bullet_list(
        doc,
        [
            "Further improve answer quality for fallback-heavy and partially answered procurement queries.",
            "Extend UAT validation with stricter semantic acceptance checks for multilingual questions.",
            "Continue frontend verification in live conditions and refine the source-viewing and answer-presentation experience.",
            "Prepare the chatbot stack and documentation for broader stakeholder review and deployment-readiness discussion.",
        ],
    )


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_base_styles(doc)
    add_header_footer(doc)
    add_title_page(doc)
    add_main_sections(doc)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
